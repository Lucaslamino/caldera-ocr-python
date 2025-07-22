import easyocr  
import pandas as pd  
import matplotlib.pyplot as plt  
import seaborn as sns  
import numpy as np  
import os  
from docx import Document  
from docx.shared import Inches  

# FUNCIÓN 1: EXTRACCIÓN DE DATOS DESDE IMAGEN
# Esta función se encarga de leer una imagen, extraer el texto y convertirlo
# en un DataFrame de pandas limpio y listo para el análisis.

def extraer_datos_desde_imagen(ruta_imagen):
    """
    Extrae datos tabulares de una imagen, filtrando valores numéricos y descartando texto.

    Args:
        ruta_imagen (str): La ruta del archivo de imagen que contiene la tabla.

    Returns:
        pd.DataFrame: Un DataFrame de Pandas con los datos extraídos y tipados numéricamente,
                      o None si ocurre un error.
    """
    try:
        # Se inicializa el lector de OCR con soporte para inglés y español.
        reader = easyocr.Reader(['en', 'es']) 
        print(f"🔍 Leyendo imagen '{ruta_imagen}'...")
        # Se lee el texto de la imagen. detail=0 devuelve solo el texto, sin coordenadas.
        resultado_ocr = reader.readtext(ruta_imagen, detail=0)

        # Se crea una lista vacía para almacenar solo los valores que sean numéricos.
        datos_numericos = []
        # Se itera sobre cada fragmento de texto extraído por el OCR.
        for texto in resultado_ocr:
            # Se limpia el texto: se quitan espacios y se reemplaza la coma decimal por un punto.
            texto_limpio = texto.strip().replace(',', '.')
            try:
                # Se intenta convertir el texto a un número flotante (float).
                float(texto_limpio)
                # Si la conversión es exitosa, se añade a la lista de datos numéricos.
                datos_numericos.append(texto_limpio)
            except ValueError:
                # Si el texto no se puede convertir a número (ej: "Presión"), se ignora.
                pass 

        # Se comprueba si el número total de datos es divisible por 8 (el número de columnas).
        if len(datos_numericos) % 8 != 0:
            # Si no es múltiplo, podría haber un error en la extracción. Se muestra una advertencia.
            print(f"⚠️ Alerta: El número de datos numéricos ({len(datos_numericos)}) no es múltiplo de 8. "
                  "Las filas podrían estar incompletas o mal extraídas.")

        # Se agrupan los datos en filas de 8 elementos cada una.
        filas = [datos_numericos[i:i+8] for i in range(0, len(datos_numericos), 8)]

        # Se definen los nombres de las columnas para el DataFrame.
        columnas = [
            "Presión (bar)", "Temperatura (°C)", "Caudal (m³/h)",
            "Nivel de Agua (%)", "Consumo de Combustible (L/h)",
            "CO (%)", "NOx (%)", "Horas Operadas"
        ]

        # Se crea el DataFrame de pandas a partir de las filas y columnas.
        df = pd.DataFrame(filas, columns=columnas)

        # Se itera sobre cada columna del DataFrame.
        for col in df.columns:
            # Se convierte cada columna a tipo numérico. 'coerce' convierte errores en NaNs (Not a Number).
            df[col] = pd.to_numeric(df[col], errors='coerce')
        
        # Se retorna el DataFrame final y limpio.
        return df

    # Manejo de errores específicos.
    except FileNotFoundError:
        # Si la imagen no se encuentra en la ruta especificada.
        print(f"❌ Archivo de imagen no encontrado: '{ruta_imagen}'.")
        return None
    except Exception as e:
        # Para cualquier otro error inesperado durante el proceso.
        print(f"❌ Ocurrió un error inesperado durante la extracción de datos: {e}")
        return None


# FUNCIÓN 2: GRÁFICO DE ANÁLISIS TEMPORAL
# Crea gráficos de líneas para visualizar cómo cambian las emisiones y el
# consumo de combustible a lo largo del tiempo (horas operadas).

def graficar_analisis_temporal(df, ruta_salida):
    """
    Genera y guarda un gráfico de líneas mostrando la evolución de emisiones (CO, NOx)
    y consumo de combustible a lo largo de las horas operadas.

    Args:
        df (pd.DataFrame): DataFrame con los datos a graficar.
        ruta_salida (str): Ruta completa para guardar la imagen del gráfico.
    """
    # Se limpian los nombres de las columnas por si tuvieran espacios extra.
    df.columns = df.columns.str.strip()

    # Se definen los nombres de las columnas a usar para mayor claridad en el código.
    eje_x = "Horas Operadas"
    emisiones_co = "CO (%)"
    emisiones_nox = "NOx (%)"
    consumo_combustible = "Consumo de Combustible (L/h)"

    # Se crea una figura para el gráfico con un tamaño específico.
    plt.figure(figsize=(14, 6))

    # Se crea el primer subgráfico (1 fila, 2 columnas, posición 1).
    ax1 = plt.subplot(1, 2, 1)
    # Se dibuja la línea para las emisiones de CO en el eje Y primario (ax1).
    sns.lineplot(x=df[eje_x], y=df[emisiones_co], marker='o', label='CO (%)', ax=ax1, color='red')
    # Se crea un segundo eje Y que comparte el mismo eje X.
    ax2 = ax1.twinx()
    # Se dibuja la línea para las emisiones de NOx en el eje Y secundario (ax2).
    sns.lineplot(x=df[eje_x], y=df[emisiones_nox], marker='o', label='NOx (%)', ax=ax2, color='blue')
    
    # Se configuran títulos y etiquetas para el primer subgráfico.
    ax1.set_title("Evolución de Emisiones")
    ax1.set_xlabel("Horas Operadas")
    ax1.set_ylabel("CO (%)", color='red')
    ax2.set_ylabel("NOx (%)", color='blue')
    ax1.grid(True) # Se añade una grilla.
    ax1.legend(loc='upper left') # Se muestra la leyenda del eje 1.
    ax2.legend(loc='upper right') # Se muestra la leyenda del eje 2.

    # Se crea el segundo subgráfico (1 fila, 2 columnas, posición 2).
    plt.subplot(1, 2, 2)
    # Se dibuja la línea para el consumo de combustible.
    sns.lineplot(x=df[eje_x], y=df[consumo_combustible], marker='o', color='green')
    
    # Se configuran títulos y etiquetas para el segundo subgráfico.
    plt.title("Evolución del Consumo de Combustible")
    plt.xlabel("Horas Operadas")
    plt.ylabel("Consumo de Combustible (L/h)")
    plt.grid(True) # Se añade una grilla.

    # Se ajusta el diseño para que los gráficos no se superpongan.
    plt.tight_layout()
    # Se guarda la figura completa en la ruta especificada.
    plt.savefig(ruta_salida)
    # Se cierra la figura para liberar memoria.
    plt.close()


# FUNCIÓN 3: GRÁFICO DE EFICIENCIA DE COMBUSTIÓN
# Genera diagramas de caja (box plots) para analizar el consumo y las
# emisiones en diferentes rangos de caudal.

def graficar_eficiencia_combustion_boxplot(df, ruta_salida):
    """
    Genera y guarda box plots para analizar el consumo de combustible y las emisiones de CO
    en función del rango de caudal.

    Args:
        df (pd.DataFrame): DataFrame con los datos a graficar.
        ruta_salida (str): Ruta completa para guardar la imagen del gráfico.
    """
    # Se obtienen los valores mínimo y máximo del caudal.
    min_caudal = df["Caudal (m³/h)"].min()
    max_caudal = df["Caudal (m³/h)"].max()

    # Se verifica si hay suficientes datos para crear los rangos.
    if pd.isna(min_caudal) or pd.isna(max_caudal) or min_caudal == max_caudal:
        print("⚠️ No hay suficiente variación en 'Caudal (m³/h)' para crear rangos. Saltando este gráfico.")
        return # Si no hay variación, se termina la función.

    # Se crean 5 rangos (bins) de igual tamaño entre el caudal mínimo y máximo.
    bins = np.linspace(min_caudal, max_caudal, 6)
    # Se crean etiquetas para cada rango (ej: "10.0 - 12.5").
    labels = [f"{round(bins[i],2)} - {round(bins[i+1],2)}" for i in range(len(bins)-1)]
    # Se hace una copia del DataFrame para no modificar el original.
    df_temp = df.copy() 
    # Se crea una nueva columna 'Caudal_Rango' que clasifica cada valor de caudal en un rango.
    df_temp['Caudal_Rango'] = pd.cut(df_temp["Caudal (m³/h)"], bins=bins, labels=labels, include_lowest=True)

    # Se crea la figura para los gráficos.
    plt.figure(figsize=(14,6))

    # Se crea el primer subgráfico (1 fila, 2 columnas, posición 1).
    plt.subplot(1, 2, 1)
    # Se genera un box plot del consumo de combustible para cada rango de caudal.
    sns.boxplot(x='Caudal_Rango', y='Consumo de Combustible (L/h)', data=df_temp, palette='Greens')
    # Se añaden títulos y etiquetas.
    plt.title('Consumo de Combustible por Rango de Caudal')
    plt.xlabel('Rango de Caudal (m³/h)')
    plt.ylabel('Consumo de Combustible (L/h)')
    plt.xticks(rotation=45) # Se rotan las etiquetas del eje X para que no se superpongan.

    # Se crea el segundo subgráfico (1 fila, 2 columnas, posición 2).
    plt.subplot(1, 2, 2)
    # Se genera un box plot de las emisiones de CO para cada rango de caudal.
    sns.boxplot(x='Caudal_Rango', y='CO (%)', data=df_temp, palette='Reds')
    # Se añaden títulos y etiquetas.
    plt.title('CO (%) por Rango de Caudal')
    plt.xlabel('Rango de Caudal (m³/h)')
    plt.ylabel('CO (%)')
    plt.xticks(rotation=45) # Se rotan las etiquetas del eje X.

    # Se ajusta el diseño y se guarda la figura.
    plt.tight_layout()
    plt.savefig(ruta_salida)
    plt.close() # Se cierra la figura.


# FUNCIÓN 4: GRÁFICO DE ANÁLISIS DE CORRELACIÓN
# Crea un gráfico de dispersión y un mapa de calor para ver la relación
# entre las diferentes variables.

def graficar_analisis_correlacion(df, ruta_salida):
    """
    Genera y guarda un scatter plot de Temperatura vs Presión y un mapa de calor de correlaciones.

    Args:
        df (pd.DataFrame): DataFrame con los datos a graficar.
        ruta_salida (str): Ruta completa para guardar la imagen del gráfico.
    """
    # Se definen los nombres de las columnas para mayor claridad.
    presion = "Presión (bar)"
    temperatura = "Temperatura (°C)"

    # Se crea la figura.
    plt.figure(figsize=(14, 6))

    # Se crea el primer subgráfico: un gráfico de dispersión (scatter plot).
    ax1 = plt.subplot(1, 2, 1)
    # Se grafica la relación entre presión (eje X) y temperatura (eje Y).
    sns.scatterplot(x=df[presion], y=df[temperatura], ax=ax1)
    # Se añaden títulos, etiquetas y una grilla.
    ax1.set_title("Temperatura vs Presión")
    ax1.set_xlabel("Presión (bar)")
    ax1.set_ylabel("Temperatura (°C)")
    ax1.grid(True)

    # Se crea el segundo subgráfico: un mapa de calor (heatmap).
    ax2 = plt.subplot(1, 2, 2)
    # Se seleccionan solo las columnas numéricas del DataFrame.
    df_numeric = df.select_dtypes(include=np.number)
    # Se calcula la matriz de correlación entre todas las variables.
    correlacion = df_numeric.corr()
    # Se dibuja el mapa de calor con los valores de correlación.
    sns.heatmap(correlacion, annot=True, fmt=".2f", cmap="coolwarm", center=0, ax=ax2)
    ax2.set_title("Mapa de Calor de Correlaciones")

    # Se ajusta el diseño, se guarda y se cierra la figura.
    plt.tight_layout()
    plt.savefig(ruta_salida)
    plt.close()


# FUNCIÓN 5: GRÁFICO DE DISTRIBUCIÓN Y ANOMALÍAS
# Crea histogramas para ver la distribución de los datos y box plots
# para detectar posibles valores atípicos (anomalías).

def graficar_distribucion_y_anomalias(df, ruta_salida):
    """
    Genera y guarda histogramas y box plots para visualizar la distribución y anomalías.

    Args:
        df (pd.DataFrame): DataFrame con los datos a graficar.
        ruta_salida (str): Ruta completa para guardar la imagen del gráfico.
    """
    # Se definen las variables para las que se crearán histogramas.
    variables_hist = ["Temperatura (°C)", "Presión (bar)", "CO (%)", "Consumo de Combustible (L/h)"]
    # Se definen las variables para las que se crearán box plots (generalmente para detectar anomalías).
    variables_box = ["CO (%)", "Consumo de Combustible (L/h)"]

    # Se crea una figura grande para alojar todos los subgráficos.
    plt.figure(figsize=(16, 10))

    # Se crea un bucle para generar los histogramas.
    for i, var in enumerate(variables_hist, 1):
        # Se crea un subgráfico en una grilla de 2x4. El índice 'i' determina la posición.
        plt.subplot(2, 4, i)
        # Se dibuja el histograma. kde=True añade una línea de densidad.
        sns.histplot(df[var], bins=30, kde=True, color='skyblue')
        # Se añaden títulos y etiquetas.
        plt.title(f"Histograma de {var}")
        plt.xlabel(var)
        plt.ylabel("Frecuencia")
        plt.grid(True, linestyle='--', alpha=0.5)

    # Se crea un bucle para generar los box plots. El índice empieza en 5 para continuar en la grilla.
    for i, var in enumerate(variables_box, 5):
        # Se crea un subgráfico en la grilla de 2x4.
        plt.subplot(2, 4, i)
        # Se dibuja el box plot.
        sns.boxplot(x=df[var], color='lightcoral')
        # Se añaden títulos y etiquetas.
        plt.title(f"Box Plot de {var}")
        plt.xlabel(var)
        plt.grid(True, linestyle='--', alpha=0.5)

    # Se ajusta el diseño, se guarda y se cierra.
    plt.tight_layout()
    plt.savefig(ruta_salida)
    plt.close()


# FUNCIÓN 6: GENERACIÓN DE INFORME EN WORD
# Consolida todos los gráficos y análisis en un único documento .docx.

def crear_informe_docx(df, directorio_graficos, nombre_informe="informe_caldera.docx"):
    """
    Crea un documento de Word (.docx) con títulos, explicaciones y los gráficos generados.

    Args:
        df (pd.DataFrame): El DataFrame con los datos (no se usa directamente pero es buena práctica pasarlo).
        directorio_graficos (str): El directorio donde están guardados los gráficos.
        nombre_informe (str): Nombre del archivo .docx a generar.
    """
    # Se crea un objeto Documento de Word en memoria.
    document = Document()

    # Se añade el título principal del informe. level=1 corresponde a "Título 1".
    document.add_heading('Informe de Análisis de Datos de Operación de Caldera', level=1)

    # Se añade un párrafo de introducción.
    document.add_paragraph(
        "Este informe presenta un análisis de los datos de operación de una caldera extraídos "
        "de una imagen. A continuación, se muestran visualizaciones clave para comprender "
        "el comportamiento del sistema..."
    )

    # --- Sección 1: Análisis Temporal ---
    # Se añade un subtítulo. level=2 corresponde a "Título 2".
    document.add_heading('Análisis de la Evolución Temporal', level=2)
    # Se añade un párrafo explicativo para esta sección.
    document.add_paragraph(
        "Esta sección presenta gráficos de líneas que muestran cómo evolucionan las emisiones "
        "de CO y NOx, así como el consumo de combustible, a lo largo de las horas de operación."
    )
    # Se construye la ruta completa a la imagen.
    ruta_grafico1 = os.path.join(directorio_graficos, 'analisis_temporal.png')
    # Se añade la imagen al documento, especificando el ancho.
    document.add_picture(ruta_grafico1, width=Inches(6.5))
    document.add_paragraph('Gráfico 1: Tendencias temporales de emisiones y consumo de combustible.')
    document.add_paragraph('\n') # Se añade un salto de línea para espaciar.

    # --- Sección 2: Eficiencia de Combustión ---
    document.add_heading('Análisis de la Eficiencia de Combustión por Caudal', level=2)
    document.add_paragraph(
        "En esta sección, se utilizan box plots para analizar la relación entre el caudal "
        "de operación, el consumo de combustible y las emisiones de CO."
    )
    ruta_grafico2 = os.path.join(directorio_graficos, 'analisis_eficiencia_boxplots.png')
    document.add_picture(ruta_grafico2, width=Inches(6.5))
    document.add_paragraph('Gráfico 2: Distribución del consumo y emisiones por rangos de caudal.')
    document.add_paragraph('\n')

    # --- Sección 3: Correlación ---
    document.add_heading('Análisis de Correlación entre Parámetros', level=2)
    document.add_paragraph(
        "Esta sección incluye un scatter plot para visualizar la relación entre temperatura y "
        "presión, y un mapa de calor que muestra las correlaciones entre todas las variables."
    )
    ruta_grafico3 = os.path.join(directorio_graficos, 'analisis_correlacion.png')
    document.add_picture(ruta_grafico3, width=Inches(6.5))
    document.add_paragraph('Gráfico 3: Relación temperatura-presión y mapa de calor de correlaciones.')
    document.add_paragraph('\n')

    # --- Sección 4: Distribución y Anomalías ---
    document.add_heading('Análisis de Distribución y Posibles Anomalías', level=2)
    document.add_paragraph(
        "Finalmente, se presentan histogramas para mostrar la distribución de variables "
        "clave y box plots para identificar posibles valores atípicos o anomalías."
    )
    ruta_grafico4 = os.path.join(directorio_graficos, 'distribucion_y_anomalias.png')
    document.add_picture(ruta_grafico4, width=Inches(6.5))
    document.add_paragraph('Gráfico 4: Distribución de variables y detección de anomalías.')
    document.add_paragraph('\n')

    # Se intenta guardar el documento en el disco.
    try:
        document.save(nombre_informe)
        print(f"\n✅ Informe '{nombre_informe}' generado con éxito.")
    except Exception as e:
        # Si ocurre un error al guardar (ej: permisos de escritura).
        print(f"❌ Error al guardar el informe: {e}")


# FUNCIÓN PRINCIPAL (main)
# Esta es la función que coordina todo el proceso: llama a las funciones
# en el orden correcto para ejecutar el script de principio a fin.

def main():
    """
    Función principal que coordina la extracción, análisis y generación de informes.
    """
    # Se definen los nombres de los archivos y directorios que se usarán.
    nombre_imagen = "caldera_table_50_observations.png"
    nombre_csv = "caldera_datos.csv"
    directorio_salida_graficos = "graficos_analisis"
    nombre_informe_final = "informe_analisis_caldera.docx"

    # Se comprueba si la carpeta para guardar los gráficos ya existe.
    if not os.path.exists(directorio_salida_graficos):
        # Si no existe, se crea.
        os.makedirs(directorio_salida_graficos)
        print(f"📂 Creado el directorio: '{directorio_salida_graficos}' para guardar los gráficos.")

    print("--- Iniciando extracción de datos y análisis ---")

    # Se llama a la función para extraer los datos de la imagen.
    df_caldera = extraer_datos_desde_imagen(nombre_imagen)

    # Se comprueba si la extracción fue exitosa y el DataFrame no está vacío.
    if df_caldera is not None and not df_caldera.empty:
        # Si es exitosa, se guarda el DataFrame en un archivo CSV.
        df_caldera.to_csv(nombre_csv, index=False)
        print(f"\n✅ Archivo CSV generado con éxito: '{nombre_csv}'")
        # Se muestra una vista previa de las primeras 5 filas de los datos.
        print("\nVista previa de los datos:")
        print(df_caldera.head())

        print("\n--- Generando gráficos de análisis ---")
        # Se llama a cada función de graficación, pasándole el DataFrame y la ruta de salida.
        graficar_analisis_temporal(df_caldera, os.path.join(directorio_salida_graficos, "analisis_temporal.png"))
        graficar_eficiencia_combustion_boxplot(df_caldera, os.path.join(directorio_salida_graficos, "analisis_eficiencia_boxplots.png"))
        graficar_analisis_correlacion(df_caldera, os.path.join(directorio_salida_graficos, "analisis_correlacion.png"))
        graficar_distribucion_y_anomalias(df_caldera, os.path.join(directorio_salida_graficos, "distribucion_y_anomalias.png"))
        print(f"\n✅ Todos los gráficos han sido generados y guardados en '{directorio_salida_graficos}'.")

        # Se llama a la función para crear el informe final en Word.
        print("\n--- Generando informe en Word ---")
        crear_informe_docx(df_caldera, directorio_salida_graficos, nombre_informe_final)

    else:
        # Si la extracción de datos falló, se muestra un mensaje de error.
        print("\n❌ No se pudo generar el DataFrame o está vacío. Revisa los errores anteriores.")

# PUNTO DE ENTRADA DEL SCRIPT
# Esta construcción asegura que la función `main()` solo se ejecute
# cuando el script es corrido directamente por el intérprete de Python.

if __name__ == "__main__":
    # Se llama a la función principal para iniciar el programa.
    main()